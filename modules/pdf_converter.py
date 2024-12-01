import streamlit as st
import tempfile
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
import fitz
import io
from PIL import Image

def convert_pdf_to_word():
    st.header("PDF转Word")
    with st.expander("PDF转换器", expanded=True):
        st.markdown("""
        ---
        该功能可以帮助你：
        - 将PDF文件转换为Word文档
        - 保留文本内容和基本格式
        - 支持多页PDF文档
        - 支持图片和表格转换
        """)
        
        uploaded_file = st.file_uploader("上传PDF文件", type=['pdf'], key="pdf_uploader")
        
        if uploaded_file is not None:
            try:
                # 创建临时文件来保存上传的PDF
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
                    tmp_pdf.write(uploaded_file.getvalue())
                    pdf_path = tmp_pdf.name

                # 创建Word文档
                doc = Document()
                
                # 打开PDF文件
                with fitz.open(pdf_path) as pdf_document:
                    total_pages = len(pdf_document)
                    
                    # 显示进度条
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    # 逐页处理PDF
                    for page_num in range(total_pages):
                        # 更新进度
                        progress = (page_num + 1) / total_pages
                        progress_bar.progress(progress)
                        status_text.text(f"正在处理第 {page_num + 1} 页，共 {total_pages} 页...")
                        
                        # 获取当前页面
                        page = pdf_document[page_num]
                        
                        # 提取图片
                        image_list = page.get_images()
                        for img_index, img in enumerate(image_list):
                            try:
                                # 获取图片信息
                                xref = img[0]
                                base_image = pdf_document.extract_image(xref)
                                image_bytes = base_image["image"]
                                
                                # 将图片字节转换为PIL Image
                                image = Image.open(io.BytesIO(image_bytes))
                                
                                # 保存为临时文件
                                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_img:
                                    image.save(tmp_img.name, format='PNG')
                                    
                                    # 添加图片到Word文档
                                    doc.add_picture(tmp_img.name, width=Inches(6.0))
                                    
                                    # 清理临时图片文件
                                    os.unlink(tmp_img.name)
                            except Exception as img_error:
                                st.warning(f"处理图片时出现警告：{str(img_error)}")
                                continue
                        
                        # 提取文本块，使用原始布局
                        text_dict = page.get_text("dict")
                        blocks = text_dict.get("blocks", [])
                        
                        for block in blocks:
                            if block.get("type") == 0:  # 文本块
                                for line in block.get("lines", []):
                                    for span in line.get("spans", []):
                                        text = span.get("text", "").strip()
                                        if text:
                                            font_size = span.get("size", 11)
                                            font_flags = span.get("flags", 0)
                                            
                                            # 创建新段落
                                            paragraph = doc.add_paragraph()
                                            run = paragraph.add_run(text)
                                            
                                            # 设置字体大小
                                            font = run.font
                                            font.size = Pt(font_size)
                                            
                                            # 设置字体样式
                                            if font_flags & 2**0:  # 粗体
                                                font.bold = True
                                            if font_flags & 2**1:  # 斜体
                                                font.italic = True
                                            
                                            # 设置颜色
                                            color = span.get("color")
                                            if color:
                                                r, g, b = [int(c * 255) for c in color]
                                                font.color.rgb = RGBColor(r, g, b)
                        
                        # 添加分页符（除了最后一页）
                        if page_num < total_pages - 1:
                            doc.add_page_break()
                
                # 保存Word文档到临时文件
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
                    doc.save(tmp_docx.name)
                    
                    # 创建下载按钮
                    with open(tmp_docx.name, 'rb') as f:
                        st.success("转换完成！")
                        st.download_button(
                            label="下载Word文档",
                            data=f,
                            file_name=f"{uploaded_file.name.rsplit('.', 1)[0]}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    
                    # 清理临时文件
                    try:
                        os.unlink(tmp_docx.name)
                    except:
                        pass
                
                # 清理临时PDF文件
                try:
                    os.unlink(pdf_path)
                except:
                    pass

            except Exception as e:
                st.error(f"转换过程中出现错误：{str(e)}")
